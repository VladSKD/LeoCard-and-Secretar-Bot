"""
Script to generate clean DOCX templates with placeholders in single runs.
Run once to regenerate templates.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def set_run_font(run, name='Times New Roman', size=14, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph_with_text(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=14, bold=False, space_after=0, space_before=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run_font(run, size=font_size, bold=bold)
    return p


def create_header(doc):
    """Create the addressee header block (right-aligned) with top spacing."""
    # Порожні рядки зверху перед "Проректору"
    for _ in range(6):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.line_spacing = 1.0

    lines = [
        "Проректору",
        "Національного університету",
        "«Львівська політехніка»",
        "Володимиру ЖЕЖУСІ",
    ]
    for line in lines:
        p = add_paragraph_with_text(doc, line, alignment=WD_ALIGN_PARAGRAPH.RIGHT, font_size=14)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def _remove_table_borders(table):
    """Remove all borders from a DOCX table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def create_signatures(doc):
    """Create the signature block at the bottom using a borderless table."""
    doc.add_paragraph()  # Empty spacer

    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    _remove_table_borders(table)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(10)
        row.cells[1].width = Cm(7.5)

    # Row 0: Голова профкому студентів і аспірантів | Богдан МІРІНЧИК
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Голова профкому")
    set_run_font(run)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_before = Pt(0)
    run = p2.add_run("студентів і аспірантів")
    set_run_font(run)

    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Богдан МІРІНЧИК")
    set_run_font(run)

    # Row 1: Голова Колегії студентів | Роксолана ПЕТРИШИН
    cell = table.cell(1, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Голова Колегії студентів")
    set_run_font(run)

    cell = table.cell(1, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Роксолана ПЕТРИШИН")
    set_run_font(run)


def create_onetime_template():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Set margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    create_header(doc)

    # Spacer
    doc.add_paragraph()

    # Title
    add_paragraph_with_text(doc, "КЛОПОТАННЯ", alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, bold=True, space_after=6)

    # Body text with placeholders (each placeholder in a single run)
    body_p = doc.add_paragraph()
    body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_p.paragraph_format.first_line_indent = Cm(1.25)

    text = "Просимо Вашого дозволу на використання приміщення "
    run = body_p.add_run(text)
    set_run_font(run)

    run = body_p.add_run("{{НОМЕР_АУДИТОРІЇ}}")
    set_run_font(run)

    run = body_p.add_run(" аудиторії ")
    set_run_font(run)

    run = body_p.add_run("{{НАЗВА_КОРПУСУ}}")
    set_run_font(run)

    run = body_p.add_run(" ")
    set_run_font(run)

    run = body_p.add_run("{{ДАТА}}")
    set_run_font(run)

    run = body_p.add_run(" року з ")
    set_run_font(run)

    run = body_p.add_run("{{ЧАС_ПОЧАТКУ}}")
    set_run_font(run)

    run = body_p.add_run(" год до ")
    set_run_font(run)

    run = body_p.add_run("{{ЧАС_КІНЦЯ}}")
    set_run_font(run)

    run = body_p.add_run(" год для проведення ")
    set_run_font(run)

    run = body_p.add_run("{{НАЗВА_ЗАХОДУ}}")
    set_run_font(run)

    run = body_p.add_run(".")
    set_run_font(run)

    # Responsible persons line
    resp_p = doc.add_paragraph()
    resp_p.paragraph_format.first_line_indent = Cm(1.25)
    resp_p.paragraph_format.space_before = Pt(2)

    run = resp_p.add_run("{{ВІДПОВІДАЛЬНІ_ЛЕЙБЛ}}")
    set_run_font(run)

    run = resp_p.add_run(" ")
    set_run_font(run)

    run = resp_p.add_run("{{ВІДПОВІДАЛЬНІ}}")
    set_run_font(run)

    run = resp_p.add_run(".")
    set_run_font(run)

    create_signatures(doc)

    doc.save('templates/template_onetime.docx')
    print("Created templates/template_onetime.docx")


def create_longterm_template():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Set margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    create_header(doc)

    # Spacer
    doc.add_paragraph()

    # Title
    add_paragraph_with_text(doc, "КЛОПОТАННЯ", alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14, bold=True, space_after=6)

    # Body text with placeholders
    body_p = doc.add_paragraph()
    body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_p.paragraph_format.first_line_indent = Cm(1.25)

    text = "Просимо Вашого дозволу на використання приміщення "
    run = body_p.add_run(text)
    set_run_font(run)

    run = body_p.add_run("{{НОМЕР_АУДИТОРІЇ}}")
    set_run_font(run)

    run = body_p.add_run(" аудиторії ")
    set_run_font(run)

    run = body_p.add_run("{{НАЗВА_КОРПУСУ}}")
    set_run_font(run)

    run = body_p.add_run(" ")
    set_run_font(run)

    run = body_p.add_run("{{ПЕРІОД_ТА_ЧАС}}")
    set_run_font(run)

    run = body_p.add_run(" для проведення ")
    set_run_font(run)

    run = body_p.add_run("{{НАЗВА_ЗАХОДУ}}")
    set_run_font(run)

    run = body_p.add_run(".")
    set_run_font(run)

    # Responsible persons line
    resp_p = doc.add_paragraph()
    resp_p.paragraph_format.first_line_indent = Cm(1.25)
    resp_p.paragraph_format.space_before = Pt(2)

    run = resp_p.add_run("{{ВІДПОВІДАЛЬНІ_ЛЕЙБЛ}}")
    set_run_font(run)

    run = resp_p.add_run(" ")
    set_run_font(run)

    run = resp_p.add_run("{{ВІДПОВІДАЛЬНІ}}")
    set_run_font(run)

    run = resp_p.add_run(".")
    set_run_font(run)

    create_signatures(doc)

    doc.save('templates/template_long.docx')
    print("Created templates/template_long.docx")


def _replace_in_doc(doc, replacements):
    """Replace placeholders in all paragraphs and table cells."""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)


def create_example_files():
    """Create filled example DOCX files for the 'Приклади клопотань' button."""
    import os
    os.makedirs('examples', exist_ok=True)

    # --- Example: onetime ---
    doc = Document('templates/template_onetime.docx')
    _replace_in_doc(doc, {
        '{{НОМЕР_АУДИТОРІЇ}}': '215',
        '{{НАЗВА_КОРПУСУ}}': 'навчального корпусу №4',
        '{{ДАТА}}': '25 вересня 2025',
        '{{ЧАС_ПОЧАТКУ}}': '14:50',
        '{{ЧАС_КІНЦЯ}}': '18:00',
        '{{НАЗВА_ЗАХОДУ}}': 'репетиції акустичного вечора',
        '{{ВІДПОВІДАЛЬНІ_ЛЕЙБЛ}}': 'Відповідальні за проведення:',
        '{{ВІДПОВІДАЛЬНІ}}': 'члени проєктного відділу Колегії та профкому студентів - Шкіль Анастасія та Лось Анастасія',
    })
    doc.save('examples/example_onetime.docx')
    print("Created examples/example_onetime.docx")

    # --- Example: longterm ---
    doc = Document('templates/template_long.docx')
    _replace_in_doc(doc, {
        '{{НОМЕР_АУДИТОРІЇ}}': '321',
        '{{НАЗВА_КОРПУСУ}}': 'головного навчального корпусу',
        '{{ПЕРІОД_ТА_ЧАС}}': 'щовівторка протягом осіннього семестру 2025/2026 навчального року з 16:25 год до 20:50 год',
        '{{НАЗВА_ЗАХОДУ}}': 'засідань Колегії та профбюро студентів ІКТА',
        '{{ВІДПОВІДАЛЬНІ_ЛЕЙБЛ}}': 'Відповідальний за проведення:',
        '{{ВІДПОВІДАЛЬНІ}}': 'Голова Колегії та профбюро студентів ІКТА - Ірина Грицай',
    })
    doc.save('examples/example_longterm.docx')
    print("Created examples/example_longterm.docx")

    # --- Example: exemption (built from scratch) ---
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # Top spacing
    for _ in range(6):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.line_spacing = 1.0

    # Header — Директорці
    for line in ["Директорці", "Інституту гуманітарних і соціальних наук",
                 "Національного університету", "«Львівська політехніка»", "доц. Зоряні КУНЬЧ"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        set_run_font(run)

    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("КЛОПОТАННЯ")
    set_run_font(run, bold=True)
    run.underline = True

    # Body
    body_p = doc.add_paragraph()
    body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_p.paragraph_format.first_line_indent = Cm(1.25)
    text = (
        "Просимо Вас вважати відсутність на заняттях "
        "17 листопада 2025 року з поважної причини у зв'язку з організацією та участю у "
        "заході «День студента 2025» студентки Лось Анастасії МВ-44."
    )
    run = body_p.add_run(text)
    set_run_font(run)

    create_signatures(doc)
    doc.save('examples/example_exemption.docx')
    print("Created examples/example_exemption.docx")


if __name__ == '__main__':
    create_onetime_template()
    create_longterm_template()
    create_example_files()
    print("All templates and examples regenerated successfully!")
