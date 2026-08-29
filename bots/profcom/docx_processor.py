from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import re
from datetime import datetime


def process_building_name(building_info: str) -> str:
    if not building_info:
        return ""
    cleaned = building_info.lower().strip()
    if cleaned in ('головний', 'гнк', 'головного', 'головного навчального корпусу'):
        return 'головного навчального корпусу'
    # Strip non-digit characters in case AI returns "корпус №4" or "4-й" etc.
    digits = re.sub(r'\D', '', cleaned)
    if digits:
        return f'навчального корпусу №{digits}'
    return f'навчального корпусу №{building_info.strip()}'


def _fix_kolehiya_capitalization(text: str) -> str:
    """Fix 'Голова колегії' → 'Голова Колегії' capitalization."""
    return re.sub(r'(?i)\bколегії\b', 'Колегії', text)


def process_responsible_persons(persons) -> str:
    """Process responsible persons data from AI.

    Accepts either a string or a list. If AI returns a pre-formatted string,
    use it directly. If AI returns structured data (list of dicts), join them.
    """
    if not persons:
        return "[відповідальні не вказані]"

    # If AI returned a single string, use it directly (it's already formatted)
    if isinstance(persons, str):
        return _fix_kolehiya_capitalization(persons.strip())

    formatted_parts = []
    for person_data in persons:
        if isinstance(person_data, str):
            formatted_parts.append(person_data.strip())
        elif isinstance(person_data, dict):
            role = person_data.get('role', person_data.get('position', '')).strip()
            name = person_data.get('name', '').strip()
            if role and name:
                formatted_parts.append(f"{role} - {name}")
            elif name:
                formatted_parts.append(name)
            elif role:
                formatted_parts.append(role)
            else:
                # Fallback: join all values
                val = " - ".join(filter(None, [str(v).strip() for v in person_data.values()]))
                if val:
                    formatted_parts.append(val)

    if not formatted_parts:
        return "[відповідальні не вказані]"

    return _fix_kolehiya_capitalization(", ".join(formatted_parts))


def replace_placeholder_in_paragraph(paragraph, key: str, value: str):
    """Replace a placeholder in a paragraph, handling cases where
    the placeholder is split across multiple runs.
    """
    full_text = paragraph.text
    if key not in full_text:
        return

    # First try: simple per-run replacement (fastest path)
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, value)
            return

    # If we get here, the placeholder is split across runs.
    runs = paragraph.runs
    if not runs:
        return

    run_texts = [run.text for run in runs]
    concat = ''.join(run_texts)

    idx = concat.find(key)
    if idx == -1:
        return

    char_pos = 0
    start_run_idx = None
    start_offset = None
    end_run_idx = None
    end_offset = None

    for i, text in enumerate(run_texts):
        run_start = char_pos
        run_end = char_pos + len(text)

        if start_run_idx is None and idx < run_end:
            start_run_idx = i
            start_offset = idx - run_start

        placeholder_end = idx + len(key)
        if end_run_idx is None and placeholder_end <= run_end:
            end_run_idx = i
            end_offset = placeholder_end - run_start
            break

        char_pos = run_end

    if start_run_idx is None or end_run_idx is None:
        return

    if start_run_idx == end_run_idx:
        t = run_texts[start_run_idx]
        runs[start_run_idx].text = t[:start_offset] + value + t[end_offset:]
    else:
        runs[start_run_idx].text = run_texts[start_run_idx][:start_offset] + value
        for mid in range(start_run_idx + 1, end_run_idx):
            runs[mid].text = ''
        runs[end_run_idx].text = run_texts[end_run_idx][end_offset:]


def _get_responsible_label(persons) -> str:
    """Return 'Відповідальний за проведення:' or 'Відповідальні за проведення:' based on count."""
    if isinstance(persons, list):
        is_plural = len(persons) > 1
    elif isinstance(persons, str):
        is_plural = ',' in persons or ' та ' in persons or ' і ' in persons
    else:
        is_plural = True
    return "Відповідальні за проведення:" if is_plural else "Відповідальний за проведення:"


# --- Helpers for dynamic document generation ---

def _set_font(run, name='Times New Roman', size=14, bold=False, underline=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline


def _remove_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_signatures(doc):
    doc.add_paragraph()  # spacer

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    _remove_borders(table)

    # Set table to full page width and remove cell padding
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Set width to 100% via XML
    width_xml = parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'
    )
    tblPr.append(width_xml)
    # Remove cell margins so text goes to very edges
    cell_mar_xml = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        '<w:left w:w="0" w:type="dxa"/>'
        '<w:right w:w="0" w:type="dxa"/>'
        '</w:tblCellMar>'
    )
    tblPr.append(cell_mar_xml)

    for row in table.rows:
        row.cells[0].width = Cm(10)
        row.cells[1].width = Cm(7.5)

    # Row 0 — Голова профкому студентів і аспірантів
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Голова профкому")
    _set_font(run)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    run = p2.add_run("студентів та аспірантів")
    _set_font(run)

    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Богдан МІРІНЧИК")
    _set_font(run)

    # Row 1 — empty spacer row
    for j in range(2):
        cell = table.cell(1, j)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Row 2 — Голова Колегії студентів
    cell = table.cell(2, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Голова Колегії студентів")
    _set_font(run)

    cell = table.cell(2, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Роксолана ПЕТРИШИН")
    _set_font(run)


def _init_doc():
    """Create a new document with standard styles and margins."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    return doc


def _add_top_spacing(doc, count=6):
    for _ in range(count):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.line_spacing = 1.0


def _save_petition(doc, filename=None):
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'Klopotannya_{timestamp}.docx'
    output_path = os.path.join('generated_petitions', filename)
    os.makedirs('generated_petitions', exist_ok=True)
    doc.save(output_path)
    return output_path


def _build_petition_filename(data: dict, petition_type: str) -> str:
    """Build filename: ДАТА — ТИП — ДЕТАЛІ.docx"""
    date_str = datetime.now().strftime("%d.%m.%Y")
    event_name = data.get('event_name', '').strip()
    # Trim event name for filename (max 60 chars)
    if len(event_name) > 60:
        event_name = event_name[:57] + "..."

    if petition_type == "onetime":
        room = data.get('auditorium_number', '')
        return f"{date_str} — Одноразова ауд.{room} — {event_name}.docx"
    elif petition_type == "longterm":
        room = data.get('auditorium_number', '')
        return f"{date_str} — Семестрова ауд.{room} — {event_name}.docx"
    else:  # exemption
        return f"{date_str} — Звільнення — {event_name}.docx"


# --- Exemption petition (dynamic generation) ---

def _add_dodatok(doc, students: list):
    """Add Додаток 1 page with numbered list of students."""
    from docx.enum.text import WD_BREAK
    # Page break
    p = doc.add_paragraph()
    p.runs[0].add_break(WD_BREAK.PAGE) if p.runs else p.add_run().add_break(WD_BREAK.PAGE)

    # "Додаток 1" header (right-aligned)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Додаток 1")
    _set_font(run)

    # Numbered list of students
    for i, student in enumerate(students, 1):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)
        name = student.get('name', '')
        group = student.get('group', '')
        run = sp.add_run(f"{i}. {name} {group}")
        _set_font(run)


def _create_exemption_petition(data: dict) -> str:
    doc = _init_doc()
    _add_top_spacing(doc, count=8)

    # Header — Директору/Директорці Інституту ...
    institute_name = data.get('institute_name', '')
    director_info = data.get('director_info', '')
    director_form = data.get('director_form', 'Директору')

    # Split long institute names into multiple lines (~40 chars max)
    institute_line = f"Інституту {institute_name}"
    if len(institute_line) > 40:
        # Find a good split point (after comma or space near the middle)
        words = institute_line.split()
        line1_parts = []
        line2_parts = []
        current_len = 0
        split_done = False
        for word in words:
            if not split_done and current_len + len(word) + 1 <= 40:
                line1_parts.append(word)
                current_len += len(word) + 1
            else:
                split_done = True
                line2_parts.append(word)
        institute_lines = [' '.join(line1_parts), ' '.join(line2_parts)]
    else:
        institute_lines = [institute_line]

    header_lines = [
        director_form,
        *institute_lines,
        "Національного університету",
        "«Львівська політехніка»",
        director_info,
    ]
    for line in header_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        _set_font(run)

    # Spacer
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("КЛОПОТАННЯ")
    _set_font(run, bold=True, underline=False)

    # Data
    students = data.get('students', [])
    students_inline = data.get('students_inline', '')
    event_date = data.get('event_date', '')
    event_name = data.get('event_name', '')

    # Clean date — remove trailing "року" (we add it in text)
    clean_date = re.sub(r'\s*року\s*$', '', event_date.strip()).strip()

    # Body text
    body_p = doc.add_paragraph()
    body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_p.paragraph_format.first_line_indent = Cm(1.25)

    if len(students) <= 6:
        # Small petition (<=6 students): "на заняттях ... з поважної причини"
        if len(students) <= 2 and students_inline:
            # Inline format: "... студентки Лось Анастасії МВ-44."
            text = (
                f"Просимо Вас вважати відсутність на парах "
                f"{clean_date} року з поважної причини у зв'язку з організацією та участю у "
                f"{event_name} {students_inline}."
            )
            run = body_p.add_run(text)
            _set_font(run)
        else:
            # List format (3-6 students): "... таких студентів:" + list
            text = (
                f"Просимо Вас вважати відсутність на парах "
                f"{clean_date} року з поважної причини у зв'язку з організацією та участю у "
                f"{event_name} таких студентів:"
            )
            run = body_p.add_run(text)
            _set_font(run)

            for i, student in enumerate(students):
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Cm(1.25)
                sp.paragraph_format.space_after = Pt(0)
                sp.paragraph_format.space_before = Pt(0)
                name = student.get('name', '')
                group = student.get('group', '')
                line_text = f"{name} {group}"
                if i == len(students) - 1:
                    line_text += "."
                run = sp.add_run(line_text)
                _set_font(run)
    else:
        # Large petition (>6 students): reference Додаток 1
        text = (
            f"Просимо Вас вважати відсутність на парах студентів, "
            f"зазначених у Додатку 1, як з поважних причин "
            f"{clean_date} року у зв'язку з організацією та участю у "
            f"{event_name}."
        )
        run = body_p.add_run(text)
        _set_font(run)

    # NO "Відповідальні" section for exemptions

    _add_signatures(doc)

    # Додаток 1 (separate page for >6 students)
    if len(students) > 6:
        _add_dodatok(doc, students)

    filename = _build_petition_filename(data, "exemption")
    return _save_petition(doc, filename)


# --- Main entry point ---

TEMPLATE_TO_TYPE = {
    "template_onetime.docx": "onetime",
    "template_long.docx": "longterm",
    "template_exemption.docx": "exemption",
}


def create_petition(template_name: str, data: dict) -> str:
    petition_type = TEMPLATE_TO_TYPE.get(template_name, "onetime")

    # Exemption is generated dynamically (different header, no responsible persons)
    if template_name == "template_exemption.docx":
        return _create_exemption_petition(data)

    template_path = os.path.join('templates', template_name)
    doc = Document(template_path)

    # Add 2 extra spacer paragraphs at the top for the header
    body = doc.element.body
    for _ in range(2):
        spacer_p = parse_xml(
            f'<w:p {nsdecls("w")}>'
            '<w:pPr><w:spacing w:after="0" w:before="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '</w:p>'
        )
        body.insert(0, spacer_p)

    # Clean date: remove trailing "року" to avoid duplication
    if 'event_date' in data and data['event_date']:
        clean_date = data['event_date'].strip()
        clean_date = re.sub(r'\s*року\s*$', '', clean_date).strip()
        data['event_date'] = clean_date

    replacements = {
        '{{НОМЕР_АУДИТОРІЇ}}': str(data.get('auditorium_number', '[не вказано]')).strip(),
        '{{НАЗВА_КОРПУСУ}}': process_building_name(data.get('building_info', '')),
        '{{НАЗВА_ЗАХОДУ}}': str(data.get('event_name', '[не вказано]')).strip(),
        '{{ВІДПОВІДАЛЬНІ_ЛЕЙБЛ}}': _get_responsible_label(data.get('responsible_persons', [])),
        '{{ВІДПОВІДАЛЬНІ}}': process_responsible_persons(data.get('responsible_persons', [])),
        '{{ДАТА}}': str(data.get('event_date', '')).strip(),
        '{{ЧАС_ПОЧАТКУ}}': str(data.get('start_time', '')).strip(),
        '{{ЧАС_КІНЦЯ}}': str(data.get('end_time', '')).strip(),
        '{{ПЕРІОД_ТА_ЧАС}}': str(data.get('period_and_time_description', '')).strip(),
    }

    # Replace in all paragraphs (body)
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            replace_placeholder_in_paragraph(paragraph, key, value)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_placeholder_in_paragraph(paragraph, key, value)

    filename = _build_petition_filename(data, petition_type)
    return _save_petition(doc, filename)
